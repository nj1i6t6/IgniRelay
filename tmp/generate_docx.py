"""
Generate DOCX document for IgniRelay project following the competition format:
  - 一、 系統名稱
  - 二、 系統目的與範圍
  - 三、 實作規劃與詳細流程
  - 四、 系統架構設計
  - 五、 數據設計與資料結構
  - 六、 系統開發環境與工具
  - 七、 系統部署方案
  - 八、 系統測試案例設計
  - 九、 測試結果分析與結論
Then convert to PDF.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = r"C:\Users\radio\Downloads\IDE\CoReM\tmp"

# ── Color Constants ──
BLUE_HEADER = RGBColor(0x1F, 0x4E, 0x79)
GREEN_HEADER = RGBColor(0x2E, 0x75, 0x4E)
GRAY_HEADER = RGBColor(0x44, 0x54, 0x6A)

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, is_header=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.name = 'Microsoft JhengHei'
        run.font.size = Pt(10)
        if is_header:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if bg_color:
            set_cell_shading(cell, bg_color)
    return row

def create_document():
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ── Style Configuration ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Microsoft JhengHei'
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
        hs.font.color.rgb = BLUE_HEADER

    # ── Header ──
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "2025 年全國大專校院智慧創新暨跨域整合創作競賽  作品設計測試文件"
    hp.style = doc.styles['Normal']
    for run in hp.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    # Add bottom border to header
    pPr = hp._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Footer with page number ──
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char)
    run = fp.add_run()
    instr = OxmlElement('w:instrText')
    instr.text = 'PAGE'
    run._r.append(instr)
    run = fp.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char)

    # ══════════════════════════════════════════════════
    # ═  Cover Page
    # ══════════════════════════════════════════════════

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2025 年全國大專校院智慧創新暨跨域整合創作競賽")
    run.font.size = Pt(22)
    run.bold = True
    run.font.name = 'Microsoft JhengHei'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("作品設計測試文件")
    run.font.size = Pt(20)
    run.bold = True
    run.font.name = 'Microsoft JhengHei'

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("烽傳 IgniRelay — 去中心化災難通訊骨幹網路")
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = BLUE_HEADER
    run.font.name = 'Microsoft JhengHei'

    for _ in range(6):
        doc.add_paragraph()

    info_items = [
        "版本：v1.0",
        "日期：2026 年 3 月",
        "階段：Phase 0 — 概念驗證與 MVP",
    ]
    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.size = Pt(14)
        run.font.name = 'Microsoft JhengHei'

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ═  Table of Contents placeholder
    # ══════════════════════════════════════════════════

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目 錄")
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = 'Microsoft JhengHei'
    doc.add_paragraph()

    toc_items = [
        "一、 系統名稱 ................................................ 3",
        "二、 系統目的與範圍 .......................................... 3",
        "三、 實作規劃與詳細流程 ...................................... 5",
        "四、 系統架構設計 ............................................ 9",
        "五、 數據設計與資料結構 ...................................... 12",
        "六、 系統開發環境與工具 ...................................... 14",
        "七、 系統部署方案 ............................................ 15",
        "八、 系統測試案例設計 ........................................ 17",
        "九、 測試結果分析與結論 ...................................... 20",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = 'Microsoft JhengHei'

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ═  一、 系統名稱
    # ══════════════════════════════════════════════════

    doc.add_heading('一、 系統名稱', level=1)
    doc.add_paragraph('烽傳 IgniRelay — 去中心化災難通訊骨幹網路')
    doc.add_paragraph(
        '「烽火台是古代第一套 Mesh 網路。烽傳，就是把這個概念裝進每個人的口袋。」'
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # ═  二、 系統目的與範圍
    # ══════════════════════════════════════════════════

    doc.add_heading('二、 系統目的與範圍', level=1)

    doc.add_heading('1. 系統目的', level=2)
    doc.add_paragraph(
        '在地震、颱風、戰爭等極端事件導致 4G/5G、光纖、WiFi 全面癱瘓時，'
        '為全台民眾與政府提供從中央到最基層里民至少 14 天的備援通訊生命線。'
    )
    doc.add_paragraph(
        '政府防災通訊可達縣市 EOC（中央與縣市之間有衛星通訊），但無法下達鄉鎮公所、'
        '消防分隊、派出所、里長、避難所與一般民眾。烽傳 IgniRelay 填補的正是這段空白：'
        '縣市政府往下到基層的最後一哩路。'
    )

    doc.add_heading('2. 系統範圍', level=2)
    doc.add_paragraph(
        '系統採四頻互補設計（LoRa 920MHz + Wi-Fi HaLow 920-928MHz + BLE 2.4GHz + 433MHz RF），'
        '涵蓋從中央行政院到一般民眾手機的完整通訊鏈。目前 Phase 0 已完成 BLE Mesh App MVP 開發與驗證。'
    )

    # Non-functional requirements table
    doc.add_heading('(1) 非功能性需求', level=3)
    nfr_table = doc.add_table(rows=1, cols=4)
    nfr_table.style = 'Table Grid'
    nfr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = nfr_table.rows[0]
    for i, text in enumerate(['編號', '需求類別', '需求描述', '達成方式']):
        hdr.cells[i].text = text
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr.cells[i], '1F4E79')

    nfr_data = [
        ['NFR-001', '離線運作', '完全不依賴行動網路/WiFi/衛星', 'BLE Mesh 自組織 P2P 網路'],
        ['NFR-002', '續航力', '至少 14 天不中斷通訊', 'LFP 電池 + 太陽能 + UPS 備援'],
        ['NFR-003', '低門檻', '民眾無需購買任何設備', '手機安裝 App 即可接入'],
        ['NFR-004', '資料一致性', '離線各自修改，重連後一致', 'HLC + Op-based CRDT'],
        ['NFR-005', '安全性', '防偽造、防篡改', 'Ed25519 數位簽章'],
        ['NFR-006', '可擴展', '支援全台 15,935 個錨點', '四頻分層 + Zone-Based 路由'],
        ['NFR-007', '跨平台', 'Android + iOS 支援', 'Flutter 框架 / 平台感知 BLE'],
        ['NFR-008', '省電', '依電量動態降級', 'Tier 分級 (T1→T2→T3)'],
    ]
    for row_data in nfr_data:
        add_table_row(nfr_table, row_data)

    doc.add_paragraph()

    # Functional requirements table
    doc.add_heading('(2) 功能性需求', level=3)
    fr_table = doc.add_table(rows=1, cols=4)
    fr_table.style = 'Table Grid'
    fr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = fr_table.rows[0]
    for i, text in enumerate(['編號', '功能名稱', '功能描述', '實作狀態']):
        hdr.cells[i].text = text
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr.cells[i], '1F4E79')

    fr_data = [
        ['FR-001', 'BLE Mesh 通訊', '手機間 BLE GATT 連線、MTU 協商、雙向資料傳輸', '✅ 完成'],
        ['FR-002', 'Bloom Filter 同步', '交換事件 ID 列表，僅傳對方缺少的差量事件', '✅ 完成'],
        ['FR-003', 'HLC 時鐘同步', '混合邏輯時鐘防止斷電歸零，維持因果排序', '✅ 完成'],
        ['FR-004', 'SOS 求救廣播', 'SOS_YELLOW / SOS_RED 緊急求救發送', '✅ 完成'],
        ['FR-005', 'Triage Queue QoS', 'SOS_RED 搶佔所有頻寬，優先傳輸', '✅ 完成'],
        ['FR-006', '物資媒合', '供需配對 + 4-PIN 實體交割確認', '✅ 完成'],
        ['FR-007', '離線地圖', '內建全台 MBTiles 離線向量地圖', '✅ 完成'],
        ['FR-008', '危險標記', '長按地圖標記危險區域，同步至 Mesh', '✅ 完成'],
        ['FR-009', '醫療卡', '離線儲存血型/過敏原/用藥，附於 SOS', '✅ 完成'],
        ['FR-010', '身份管理', 'Trust Ladder L0-L3 漸進信任', '✅ 完成'],
        ['FR-011', '地理圍欄路由', 'Zone-Based 行政區邊界路由決策', '✅ 完成'],
        ['FR-012', '惡意節點隔離', '加權投票 Quarantine 去中心化黑名單', '✅ 完成'],
        ['FR-013', '聊天室', '基於 Event Log 的群組訊息功能', '📋 規劃中'],
        ['FR-014', '加密傳輸', 'ChaCha20-Poly1305 + X25519 E2EE', '📋 規劃中'],
        ['FR-015', '省電排程', 'Power Profile 動態掃描間隔', '📋 規劃中'],
        ['FR-016', '433MHz 火警', 'Si4467 住警器 RF 轉譯', 'Phase 1'],
    ]
    for row_data in fr_data:
        add_table_row(fr_table, row_data)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ═  三、 實作規劃與詳細流程
    # ══════════════════════════════════════════════════

    doc.add_heading('三、 實作規劃與詳細流程', level=1)

    doc.add_heading('1. BLE Mesh 通訊流程', level=2)
    doc.add_paragraph(
        '烽傳 App 的核心通訊採用 BLE (Bluetooth Low Energy) GATT 協定進行 Mesh 組網。'
        '每台手機同時擔任 Central（主動掃描連線）和 Peripheral（GATT Server 被動接收）角色。'
    )

    steps = [
        '① 啟動 BLE 掃描，過濾 IgniRelay Service UUID (a4d11949...)',
        '② 發現節點後加入待連佇列，序列化處理（避免 Android GATT 133 錯誤）',
        '③ 連線成功 → MTU 協商（目標 517 bytes）→ 服務發現',
        '④ 讀取對端 Bloom Filter Characteristic（事件 ID 列表）',
        '⑤ 比對差集，計算對方缺少的事件',
        '⑥ 優先推送 TriageQueue 中的高優先事件（SOS_RED 先行）',
        '⑦ 從 SQLite DB 補充最近 24 小時的事件',
        '⑧ 訂閱 Notify → 接收對端推送的事件',
        '⑨ 對收到的事件進行：Protobuf 解碼 → 去重 → HLC merge → Zone 路由判斷 → DB 寫入',
        '⑩ 同步完成 → 設定 60 秒冷卻時間 → 斷線 → 繼續掃描下一個節點',
    ]
    for step in steps:
        p = doc.add_paragraph(step, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_heading('2. 物資媒合流程', level=2)
    doc.add_paragraph(
        '物資媒合採用去中心化兩階段設計，所有配對邏輯由事件驅動，無需中央伺服器。'
    )
    steps = [
        '① 物資提供方發送 RESOURCE_REGISTER 事件（類型、數量、GPS、服務半徑）',
        '② 需求方發送 REQUEST_BROADCAST 事件（需求類型、數量、緊急度、位置）',
        '③ 收到雙方事件的節點執行本地配對演算（距離+緊急度+評分）',
        '④ 產生 MATCH_INTENT 事件廣播配對結果',
        '⑤ 衝突解決：HLC 先發先得 → Urgency 高者勝 → PubKey 字典序 tiebreak',
        '⑥ 雙方面對面以 4 位 PIN 碼完成 PHYSICAL_HANDSHAKE 交割',
        '⑦ PIN 驗證 3 次失敗 → 發送 MATCH_CANCEL → 釋放資源回到可配對池',
    ]
    for step in steps:
        p = doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('3. SOS 求救流程', level=2)
    doc.add_paragraph('SOS 求救分為兩級：SOS_YELLOW（非立即生命危險）和 SOS_RED（立即生命危險）。')
    steps = [
        '① 使用者長按 3 秒觸發 SOS 發送',
        '② 封包攜帶：GPS 座標、醫療卡摘要（MedicalSummary）、身份等級',
        '③ SOS_RED 進入 Triage Queue，搶佔所有一般資料傳輸頻寬',
        '④ 路由：SOS_RED（identity≥1）不受行政區邊界限制；SOS_YELLOW 限鄉鎮市區',
        '⑤ 節點收到 SOS 後寫入 DB + 顯示於地圖層（危險標記）',
    ]
    for step in steps:
        p = doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('4. 聊天室功能規劃 (後續)', level=2)
    doc.add_paragraph(
        '聊天室設計為 append-only event log，與現有 Event_Logs 完全同構。'
        '新增 CHAT_MESSAGE EventType (=8)，搭配 ChatMessageData sub-message。'
        '每則訊息以 HLC 排序，透過 Bloom Filter 差異同步。'
        '私聊功能可搭配 Per-peer E2EE (X25519 + ChaCha20-Poly1305) 加密。'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ═  四、 系統架構設計
    # ══════════════════════════════════════════════════

    doc.add_heading('四、 系統架構設計', level=1)

    doc.add_heading('1. 整體系統架構', level=2)
    doc.add_paragraph(
        '系統分為四層架構：骨幹層（LoRa + HaLow）、接入層（BLE）、感知層（433MHz RF）、'
        '以及應用層（手機 App + 管理儀表板）。'
    )

    arch_table = doc.add_table(rows=1, cols=4)
    arch_table.style = 'Table Grid'
    hdr = arch_table.rows[0]
    for i, text in enumerate(['層級', '技術', '範圍', '功能']):
        hdr.cells[i].text = text
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr.cells[i], '1F4E79')

    arch_data = [
        ['骨幹層', 'LoRa (LR2021)', '5-50 km', 'SOS 直達、HaLow 喚醒、心跳'],
        ['資料層', 'Wi-Fi HaLow (MM8108)', '0.8-3 km', '視訊通話、圖片、指揮同步'],
        ['接入層', 'BLE 6.0 (nRF54H20)', '~100 m', '手機 App、節點管理'],
        ['感知層', '433MHz RF (Si4467)', '建築物穿透', '住警器火警訊號接收'],
    ]
    for row_data in arch_data:
        add_table_row(arch_table, row_data)

    doc.add_paragraph()

    doc.add_heading('2. App 軟體架構 (lib/ 目錄)', level=2)

    app_table = doc.add_table(rows=1, cols=3)
    app_table.style = 'Table Grid'
    hdr = app_table.rows[0]
    for i, text in enumerate(['模組', '檔案', '職責']):
        hdr.cells[i].text = text
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr.cells[i], '2E754E')

    app_data = [
        ['BLE 通訊', 'mesh/ble_manager.dart (774 行)', 'Central 角色掃描+連線+同步'],
        ['事件管理', 'mesh/event_manager.dart', 'EventManager + TriageQueue'],
        ['事件處理', 'mesh/mesh_event_handler.dart (356 行)', '解碼、去重、HLC merge、DB 寫入'],
        ['序列化', 'mesh/event_serializer.dart (275 行)', 'Protobuf 序列化/反序列化'],
        ['路由', 'mesh/mesh_router.dart (113 行)', 'Zone-Based Geo-Fencing'],
        ['原生橋接', 'mesh/native_bridge.dart', 'MethodChannel ↔ Android Nordic'],
        ['CRDT', 'crdt/hlc.dart (89 行)', 'HLC 混合邏輯時鐘'],
        ['衝突解決', 'crdt/conflict_resolver.dart (40 行)', 'Double-Spending 排解'],
        ['加密', 'crypto/identity_manager.dart (124 行)', 'Ed25519 金鑰 + Trust Ladder'],
        ['Protobuf', 'proto/mesh_protocol.pb.dart', '自動生成的 Protobuf 綁定'],
        ['資料庫', 'db/database_helper.dart', 'SQLite 資料庫管理'],
        ['地理', 'geo/village_geofence.dart', '行政區邊界查詢'],
    ]
    for row_data in app_data:
        add_table_row(app_table, row_data)

    doc.add_paragraph()

    doc.add_heading('3. Protobuf 協議設計', level=2)
    doc.add_paragraph(
        '系統使用 Protocol Buffers 3 作為序列化格式，定義於 protos/mesh_protocol.proto (189 行)。'
        '核心封包 MeshEvent 包含事件 ID、發送者公鑰、HLC 時鐘、TTL、GPS 座標、payload 與簽章。'
    )

    proto_table = doc.add_table(rows=1, cols=3)
    proto_table.style = 'Table Grid'
    hdr = proto_table.rows[0]
    for i, text in enumerate(['EventType', '值', '說明']):
        hdr.cells[i].text = text
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr.cells[i], '44546A')

    proto_data = [
        ['RESOURCE_REGISTER', '0', '物資登記（ID/類型/數量/位置/服務半徑）'],
        ['REQUEST_BROADCAST', '1', '需求廣播（類型/數量/緊急度/位置）'],
        ['MATCH_INTENT', '2', '意向匹配（雙方公鑰/評分/超時）'],
        ['PHYSICAL_HANDSHAKE', '3', '物理核銷交割（雙方簽章/交割方式）'],
        ['HAZARD_MARKER', '4', '動態危險標記（類型/嚴重度/位置/半徑）'],
        ['QUARANTINE_VOTE', '5', '惡意節點檢舉投票（目標公鑰/權重）'],
        ['MATCH_CANCEL', '6', '釋放配對（取消原因）'],
        ['FIRE_ALARM_RF', '7', '住警器火警訊號（品牌/RSSI/座標）'],
        ['CHAT_MESSAGE', '8', '聊天訊息（規劃中）'],
    ]
    for row_data in proto_data:
        add_table_row(proto_table, row_data)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ═  五、 數據設計與資料結構
    # ══════════════════════════════════════════════════

    doc.add_heading('五、 數據設計與資料結構', level=1)

    doc.add_heading('1. Event_Logs 表', level=2)
    doc.add_paragraph('核心事件日誌表，所有 Mesh 事件（SOS、物資、危險標記等）均存入此表。')

    ev_table = doc.add_table(rows=1, cols=4)
    ev_table.style = 'Table Grid'
    hdr = ev_table.rows[0]
    for i, text in enumerate(['欄位', '型態', '說明', '備註']):
        hdr.cells[i].text = text
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr.cells[i], '44546A')

    ev_data = [
        ['event_id', 'TEXT PRIMARY KEY', '事件唯一 ID', 'UUID v4'],
        ['sender_pub_key', 'BLOB', 'Ed25519 公鑰', '32 bytes'],
        ['identity_level', 'INTEGER', '信任等級 0-3', ''],
        ['event_type', 'INTEGER', 'EventType 枚舉值', '0-7'],
        ['urgency', 'INTEGER', 'UrgencyLevel 枚舉值', '0-3'],
        ['hlc_timestamp', 'INTEGER', 'HLC 物理時間 (Unix ms)', ''],
        ['hlc_counter', 'INTEGER', 'HLC 邏輯計數器', ''],
        ['ttl', 'INTEGER', '剩餘存活跳數', '預設 10'],
        ['received_lat', 'REAL', '接收者緯度', '每跳覆寫'],
        ['received_lng', 'REAL', '接收者經度', '每跳覆寫'],
        ['origin_lat', 'REAL', '創建者緯度', '不可變'],
        ['origin_lng', 'REAL', '創建者經度', '不可變'],
        ['payload', 'BLOB', '序列化業務資料', 'Protobuf'],
        ['signature', 'BLOB', 'Ed25519 簽章', '64 bytes'],
        ['is_synced', 'INTEGER', '是否已同步', '0/1'],
    ]
    for row_data in ev_data:
        add_table_row(ev_table, row_data)

    doc.add_paragraph()

    doc.add_heading('2. Hazards_State 表', level=2)
    doc.add_paragraph('危險區域狀態表，由 HAZARD_MARKER 事件觸發寫入。')

    hz_table = doc.add_table(rows=1, cols=3)
    hz_table.style = 'Table Grid'
    hdr = hz_table.rows[0]
    for i, text in enumerate(['欄位', '型態', '說明']):
        hdr.cells[i].text = text
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr.cells[i], '44546A')

    hz_data = [
        ['hazard_id', 'TEXT PRIMARY KEY', '危險標記唯一 ID'],
        ['type', 'TEXT', '類型 (FIRE/FLOOD/ROADBLOCK/CHEMICAL)'],
        ['severity', 'INTEGER', '嚴重度 1-5'],
        ['lat / lng', 'REAL', '中心點座標'],
        ['radius', 'REAL', '影響半徑 (公尺)'],
        ['reported_by', 'TEXT', '回報者公鑰 hex'],
        ['confirm_count', 'INTEGER', '確認次數'],
    ]
    for row_data in hz_data:
        add_table_row(hz_table, row_data)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ═  六、 系統開發環境與工具
    # ══════════════════════════════════════════════════

    doc.add_heading('六、 系統開發環境與工具', level=1)

    doc.add_heading('1. 開發技術棧', level=2)

    dev_table = doc.add_table(rows=1, cols=3)
    dev_table.style = 'Table Grid'
    hdr = dev_table.rows[0]
    for i, text in enumerate(['類別', '技術/工具', '版本/說明']):
        hdr.cells[i].text = text
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr.cells[i], '2E754E')

    dev_data = [
        ['App 框架', 'Flutter', 'SDK ≥3.2.0 <4.0.0'],
        ['程式語言', 'Dart', '3.2+'],
        ['BLE (iOS)', 'flutter_blue_plus', '^2.2.1'],
        ['BLE (Android)', 'Nordic BLE Library', 'via MethodChannel (Kotlin)'],
        ['序列化', 'Protocol Buffers 3', 'protobuf ^3.1.0 / protoc_plugin ^21.1.1'],
        ['加密', 'cryptography (Dart)', '^2.7.0 (Ed25519)'],
        ['地圖', 'flutter_map + vector_map_tiles', '^7.0.2 / ^8.0.0'],
        ['離線地圖', 'mbtiles', '^0.4.2 (OpenMapTiles schema)'],
        ['資料庫', 'sqflite + sqlite3', '^2.3.3 / ^2.9.0'],
        ['狀態管理', 'Provider', '^6.1.2'],
        ['定位', 'geolocator', '^11.1.0'],
        ['權限', 'permission_handler', '^11.3.1'],
        ['IDE', 'Android Studio / VS Code', ''],
        ['版控', 'Git', ''],
        ['測試', 'flutter_test + sqflite_common_ffi', ''],
    ]
    for row_data in dev_data:
        add_table_row(dev_table, row_data)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ═  七、 系統部署方案
    # ══════════════════════════════════════════════════

    doc.add_heading('七、 系統部署方案', level=1)

    doc.add_heading('1. App 部署', level=2)
    doc.add_paragraph('手機端 App 以 APK/IPA 形式發布，民眾直接安裝即可使用。')

    app_deploy = doc.add_table(rows=1, cols=3)
    app_deploy.style = 'Table Grid'
    hdr = app_deploy.rows[0]
    for i, text in enumerate(['項目', '規格', '說明']):
        hdr.cells[i].text = text
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr.cells[i], '44546A')

    ad_data = [
        ['APK 大小', '~260 MB', '含全台離線向量地圖'],
        ['最低 Android', 'API 23 (Android 6.0)', 'BLE 5.0 建議 API 26+'],
        ['最低 iOS', 'iOS 13.0+', 'CoreBluetooth BLE'],
        ['權限需求', 'Bluetooth + Location', '背景 BLE 需前景服務'],
        ['離線能力', '100% 離線運作', '不需要任何網路連線'],
    ]
    for row_data in ad_data:
        add_table_row(app_deploy, row_data)

    doc.add_paragraph()

    doc.add_heading('2. 硬體節點部署', level=2)
    doc.add_paragraph(
        '全台部署規模估算: 約 15,935 個錨點節點（含 368 公所、628 消防分隊、1,491 派出所、'
        '5,928 避難所、102 醫院、7,752 里辦公室等）。分三階段建置。'
    )

    deploy_table = doc.add_table(rows=1, cols=4)
    deploy_table.style = 'Table Grid'
    hdr = deploy_table.rows[0]
    for i, text in enumerate(['階段', '範圍', '錨點數', '硬體成本']):
        hdr.cells[i].text = text
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr.cells[i], '44546A')

    dp_data = [
        ['Phase 0 (現況)', 'BLE Mesh App MVP', '0', 'NT$0'],
        ['Phase 1 (鄉鎮試點)', '單一偏遠鄉鎮', '~35', 'NT$50-80 萬'],
        ['Phase 2 (縣市試點)', '單縣市關鍵設施', '~650', 'NT$800-1,300 萬'],
        ['Phase 3 (全台)', '22 縣市', '~8,000+', '待實測後修正'],
    ]
    for row_data in dp_data:
        add_table_row(deploy_table, row_data)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ═  八、 系統測試案例設計
    # ══════════════════════════════════════════════════

    doc.add_heading('八、 系統測試案例設計', level=1)

    test_cases = [
        ['TC-001', 'BLE 掃描發現', '開啟兩台手機 App', '兩台手機在 BLE 範圍內', 
         '掃描到對方 IgniRelay Service UUID', '✅ Pass'],
        ['TC-002', 'MTU 協商', '連線後自動協商 MTU', 'BLE 5.0+ 設備',
         'MTU 達到 517 bytes', '✅ Pass'],
        ['TC-003', 'Bloom Filter 交換', '讀取對端 Bloom Char', '兩端各有事件',
         '正確識別差異事件 ID', '✅ Pass'],
        ['TC-004', '事件同步', '發送 RESOURCE 事件', '手機 A 創建事件',
         '手機 B 收到並存入 DB', '✅ Pass'],
        ['TC-005', 'HLC 合併', '兩台設備不同時間', '系統時鐘差異 >1 秒',
         'HLC merge 後時間不倒退', '✅ Pass'],
        ['TC-006', '去重 (SKIP seen)', '同一事件發送兩次', '重複事件 ID',
         '第二次被跳過不重複寫入', '✅ Pass'],
        ['TC-007', 'SOS 搶佔', '同時有一般事件和 SOS', 'TriageQueue 有混合事件',
         'SOS_RED 先被傳輸', '✅ Pass'],
        ['TC-008', '地理圍欄路由', '不同里的事件', '事件 origin 跨里',
         'INFO 事件不被跨里轉發', '✅ Pass'],
        ['TC-009', '物資配對', '建立 RESOURCE + REQUEST', '兩者在同一里',
         '產生 MATCH_INTENT', '✅ Pass'],
        ['TC-010', '跨廠牌 BLE', 'Pixel ↔ OPPO/Samsung', '不同廠牌 Android',
         '成功連線並同步', '✅ Pass'],
        ['TC-011', '30 秒 Timeout', '連線到無回應設備', '對端 GATT 不回應',
         '30 秒後取消並繼續下一個', '✅ Pass'],
        ['TC-012', 'OPPO Blind Relay', 'OPPO 設備 GATT read 失敗', 'ColorOS BLE 問題',
         '進入 blind relay 模式發送全部事件', '✅ Pass'],
    ]

    tc_table = doc.add_table(rows=1, cols=6)
    tc_table.style = 'Table Grid'
    hdr = tc_table.rows[0]
    for i, text in enumerate(['編號', '測試項目', '操作步驟', '前置條件', '預期結果', '結果']):
        hdr.cells[i].text = text
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr.cells[i], '1F4E79')

    for tc in test_cases:
        add_table_row(tc_table, tc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ═  九、 測試結果分析與結論
    # ══════════════════════════════════════════════════

    doc.add_heading('九、 測試結果分析與結論', level=1)

    doc.add_heading('1. 測試結果摘要', level=2)
    doc.add_paragraph(
        '所有 12 項核心測試案例均通過。BLE Mesh 通訊在跨廠牌 Android 設備上驗證成功，'
        'Bloom Filter 差異同步、HLC 時鐘合併、Triage Queue QoS 搶佔、Zone-Based 路由等'
        '核心演算法均如預期運作。'
    )

    result_table = doc.add_table(rows=1, cols=3)
    result_table.style = 'Table Grid'
    hdr = result_table.rows[0]
    for i, text in enumerate(['模組', '測試結果', '驗證方式']):
        hdr.cells[i].text = text
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr.cells[i], '2E754E')

    result_data = [
        ['BLE Mesh 多節點通訊', '✅ 通過', '實機跨廠牌測試 + debug log'],
        ['Bloom Filter 差異同步', '✅ 通過', '實機 log 分析 (大小正確縮放)'],
        ['HLC + CRDT 離線同步', '✅ 通過', '單元測試 + 整合測試'],
        ['Triage Queue QoS', '✅ 通過', '模擬測試'],
        ['地理圍欄路由', '✅ 通過', '程式邏輯驗證'],
        ['物資媒合 + 4-PIN', '✅ 通過', '功能測試'],
        ['Ed25519 簽章', '✅ 通過', '加密套件驗證'],
        ['SQLite 持久化', '✅ 通過', '重啟後資料保留'],
    ]
    for row_data in result_data:
        add_table_row(result_table, row_data)

    doc.add_paragraph()

    doc.add_heading('2. 已知限制與後續改進', level=2)
    limitations = [
        '省電：目前掃描 duty cycle ~86%，計畫導入 Power Profile 動態排程 (duty < 5%)',
        '加密：目前僅有簽章驗證，計畫加入 ChaCha20-Poly1305 傳輸層加密',
        '聊天：計畫新增 CHAT_MESSAGE EventType，基於既有 event log 架構',
        'Coded PHY：BLE 5.0 長距離模式 (300m+) 有待 flutter_blue_plus 支援完善',
        'iOS 背景限制：CoreBluetooth 背景模式有 10 分鐘限制，骨幹以實體節點為主',
    ]
    for item in limitations:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3. 結論', level=2)
    doc.add_paragraph(
        '烽傳 IgniRelay 在 Phase 0 階段已完成 BLE Mesh 核心通訊層的實作與驗證。'
        '所有核心演算法（HLC、CRDT、Bloom Filter、Triage Queue、Geo-Fencing、Quarantine）'
        '均已實作並通過測試。系統設計對齊台灣六層行政指揮體系，技術路徑與 HaLert 論文 '
        '(2025, arXiv 2507.07841) 的學術驗證一致。Phase 1 將進行 LoRa + HaLow + BLE 三頻'
        '骨幹硬體整合，取得台灣地形下的真實覆蓋率、功耗與延遲數據。'
    )

    # ── Save DOCX ──
    output_path = os.path.join(OUTPUT_DIR, "烽傳_IgniRelay_作品設計測試文件.docx")
    doc.save(output_path)
    print(f"[OK] DOCX saved: {output_path}")
    return output_path


if __name__ == "__main__":
    docx_path = create_document()
    print(f"\nDone! File saved to: {docx_path}")
